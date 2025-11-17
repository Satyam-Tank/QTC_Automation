import os
import pandas as pd
import PyPDF2
from docx import Document
from typing import Optional, Dict, List, Any

class DocumentProcessor:
    """Process various document formats (Excel, PDF, Word) to extract text."""

    def __init__(self) -> None:
        self.supported_formats = {
            '.xlsx': self.process_excel,
            '.xls': self.process_excel,
            '.pdf': self.process_pdf,
            '.docx': self.process_word,
            '.doc': self.process_word,
            '.jpg': 'image',
            '.jpeg': 'image',
            '.png': 'image'
        }

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF file."""
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""

    def extract_text_from_word(self, docx_path: str) -> str:
        """Extract text from Word document."""
        try:
            doc = Document(docx_path)
            text: List[str] = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text: List[str] = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))

            return "\n".join(text)
        except Exception as e:
            print(f"Error extracting text from Word document: {e}")
            return ""

    def extract_data_from_excel(self, excel_path: str) -> List[Dict[str, Any]]:
        """Extract data from Excel file, handling multiple sheets."""
        try:
            excel_file = pd.ExcelFile(excel_path)
            all_data: List[Dict[str, Any]] = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_path, sheet_name=sheet_name)
                
                # NaNs can cause issues, fill them
                df = df.fillna("")
                
                text_data = f"Sheet: {sheet_name}\n"
                text_data += df.to_string(index=False)

                sheet_data = {
                    'sheet_name': sheet_name,
                    'text': text_data,
                    'dataframe': df,
                    'rows': len(df)
                }
                all_data.append(sheet_data)

            return all_data
        except Exception as e:
            print(f"Error extracting data from Excel: {e}")
            return []

    def process_excel(self, file_path: str) -> Dict[str, Any]:
        """Process Excel file and return structured data."""
        excel_data = self.extract_data_from_excel(file_path)
        combined_text = "\n\n".join([sheet['text'] for sheet in excel_data])

        return {
            'type': 'excel',
            'text': combined_text,
            'sheets': excel_data,
            'sheet_count': len(excel_data)
        }

    def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF file and return extracted text."""
        text = self.extract_text_from_pdf(file_path)
        return {'type': 'pdf', 'text': text}

    def process_word(self, file_path: str) -> Dict[str, Any]:
        """Process Word document and return extracted text."""
        text = self.extract_text_from_word(file_path)
        return {'type': 'word', 'text': text}

    def identify_document_type(self, file_path: str) -> str:
        """Identify the type of document based on extension."""
        ext = os.path.splitext(file_path)[1].lower()
        return self.supported_formats.get(ext, 'unknown')

    def process_document(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Process any supported document format."""
        doc_type = self.identify_document_type(file_path)

        if doc_type == 'unknown':
            print(f"Unsupported document type for file: {file_path}")
            return None
        elif doc_type == 'image':
            return {'type': 'image', 'path': file_path}
        
        processor_func = self.supported_formats.get(os.path.splitext(file_path)[1].lower())
        
        if callable(processor_func):
            return processor_func(file_path)
        
        return None
